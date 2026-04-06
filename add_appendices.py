#!/usr/bin/env python3
"""
1. Добавляет поясняющие комментарии к кодовым фрагментам внутри отчёта.
2. Добавляет Приложение А (main.c) и Приложение Б (server.c) в конец документа.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

DOC = "Практическая_работа_HTTP_сервер.docx"

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def para_index(doc, fragment):
    for i, p in enumerate(doc.paragraphs):
        if fragment in p.text:
            return i
    return -1

def set_mono(para, size_pt=8.5):
    """Monospaced style for code paragraphs."""
    for run in para.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(size_pt)
        run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)
    pPr = para._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EDF2F4")
    pPr.append(shd)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:right"), "180")
    pPr.append(ind)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), "40")
    sp.set(qn("w:after"),  "40")
    sp.set(qn("w:line"),   "240")
    pPr.append(sp)
    bdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "B0C4D8")
        bdr.append(b)
    pPr.append(bdr)

def add_para_after(doc, ref_idx, text, style="Normal", mono=False,
                   bold=False, italic=False, size=None, color=None,
                   align=None):
    """Insert paragraph after ref_idx; return new paragraph."""
    ref = doc.paragraphs[ref_idx]._element
    p_el = OxmlElement("w:p")
    ref.addnext(p_el)
    p = Paragraph(p_el, doc)
    if style != "Normal":
        try:
            p.style = doc.styles[style]
        except Exception:
            pass
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        if bold:
            run.font.bold = bold
        if italic:
            run.font.italic = italic
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    if mono:
        set_mono(p)
    return p

def add_heading(doc, ref_idx, text, level=1):
    """Insert a heading paragraph after ref_idx."""
    ref = doc.paragraphs[ref_idx]._element
    p_el = OxmlElement("w:p")
    ref.addnext(p_el)
    p = Paragraph(p_el, doc)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = RGBColor(0x1D, 0x35, 0x57)
    pPr = p_el.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), "200")
    sp.set(qn("w:after"),  "80")
    pPr.append(sp)
    return p

def add_code_block(doc, ref_idx, code_text, comment_text=""):
    """
    Insert: optional comment paragraph + code paragraph after ref_idx.
    Returns index of last inserted paragraph.
    """
    idx = ref_idx
    if comment_text:
        p = add_para_after(doc, idx, comment_text, italic=True,
                           color=RGBColor(0x44, 0x44, 0x55), size=9.5)
        pPr = p._element.get_or_add_pPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), "80")
        sp.set(qn("w:after"),  "20")
        pPr.append(sp)
        # find new index
        for i, para in enumerate(doc.paragraphs):
            if para._element is p._element:
                idx = i
                break
    code_p = add_para_after(doc, idx, code_text, mono=True)
    for i, para in enumerate(doc.paragraphs):
        if para._element is code_p._element:
            return i
    return idx + 1

def last_para_index(doc):
    return len(doc.paragraphs) - 1

# ─────────────────────────────────────────────────────────────────────────────
# Inline comment patches
# ─────────────────────────────────────────────────────────────────────────────

INLINE_PATCHES = [
    # (фрагмент параграфа перед которым ставим блок,
    #  текст комментария,
    #  вставить ПОСЛЕ этого параграфа)
    (
        "int s = socket(AF_INET, SOCK_STREAM, 0);",
        "Листинг 1 — Создание TCP-сокета и привязка к адресу. SO_REUSEADDR позволяет "
        "немедленно переиспользовать порт после перезапуска сервера.",
    ),
    (
        "struct timeval tv;",
        "Листинг 2 — Установка тайм-аутов сокета. Защищает от зависания на медленных "
        "клиентах: если за 10 с данные не пришли — соединение закрывается.",
    ),
    (
        'if (strcmp(req.method, "GET") == 0)',
        "Листинг 3 — Маршрутизация запросов. Сначала проверяется GET (форма не требует "
        "авторизации), затем POST с проверкой учётных данных.",
    ),
    (
        'sscanf(buf, "%15s %255s", req->method, req->path);',
        "Листинг 4 — Парсинг строки запроса и заголовков. Ограничители ширины (%7s, "
        "%127s) предотвращают переполнение буфера; strstr находит нужный заголовок за O(n).",
    ),
    (
        "static int b64val(char c)",
        "Листинг 5 — Base64-декодер. Каждый символ несёт 6 бит; "
        "биты накапливаются в аккумуляторе v и извлекаются побайтово.",
    ),
    (
        "typedef struct {",
        "Листинг 6 — Структура конфигурации. Все поля хранятся как массивы фиксированного "
        "размера — никакого malloc, структура копируется по значению.",
    ),
    (
        'char key[64], value[256];',
        "Листинг 7 — Парсер конфигурационного файла. Формат key=value; "
        "строки без '=' молча пропускаются, что допускает комментарии.",
    ),
]

# ─────────────────────────────────────────────────────────────────────────────
# Source files for appendices
# ─────────────────────────────────────────────────────────────────────────────

def read_src(path):
    with open(path, "r") as f:
        return f.read().strip()

SRC = {
    "main.c":      read_src("src/main.c"),
    "server.c":    read_src("src/server.c"),
    "server.h":    read_src("src/server.h"),
    "multipart.c": read_src("src/multipart.c"),
    "auth.c":      read_src("src/auth.c"),
    "config.c":    read_src("src/config.c"),
}

# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    doc = Document(DOC)

    # ── 1. Insert inline code comments ───────────────────────────────────────
    # Work top-to-bottom; after each insert re-find anchor since indices shift
    inserted = 0
    for anchor, comment in INLINE_PATCHES:
        idx = para_index(doc, anchor)
        if idx < 0:
            print(f"  SKIP (not found): {anchor[:50]}")
            continue
        # Insert italic comment BEFORE the code paragraph
        # i.e. after idx-1
        insert_after = idx - 1
        if insert_after < 0:
            insert_after = 0
        ref = doc.paragraphs[insert_after]._element
        p_el = OxmlElement("w:p")
        ref.addnext(p_el)
        p = Paragraph(p_el, doc)
        run = p.add_run(comment)
        run.font.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x33, 0x44, 0x66)
        pPr = p_el.get_or_add_pPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), "100")
        sp.set(qn("w:after"),  "20")
        pPr.append(sp)
        inserted += 1
        print(f"  + comment before '{anchor[:45]}'")

    print(f"Inserted {inserted} inline comments")

    # ── 2. Append Приложение А and Б ─────────────────────────────────────────
    def append(text, mono=False, bold=False, size=None, color=None,
                align=None, spacing_before=0, spacing_after=0,
                page_break=False):
        p = doc.add_paragraph()
        if page_break:
            run = p.add_run()
            run.add_break(__import__("docx.enum.text", fromlist=["WD_BREAK"])
                          .WD_BREAK.PAGE)
        if text:
            run = p.add_run(text)
            run.font.bold = bold
            if size:
                run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color
            if mono:
                run.font.name = "Courier New"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)
        if align:
            p.alignment = align
        if mono:
            set_mono(p)
        pPr = p._element.get_or_add_pPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), str(spacing_before))
        sp.set(qn("w:after"),  str(spacing_after))
        pPr.append(sp)
        return p

    # ── Приложение А — main.c ─────────────────────────────────────────────────
    append("", page_break=True)
    append("ПРИЛОЖЕНИЕ А", bold=True, size=13,
           color=RGBColor(0x1D, 0x35, 0x57),
           align=WD_ALIGN_PARAGRAPH.CENTER,
           spacing_before=0, spacing_after=40)
    append("Листинг А.1 — Точка входа (main.c)",
           bold=True, size=10,
           color=RGBColor(0x1D, 0x35, 0x57),
           align=WD_ALIGN_PARAGRAPH.CENTER,
           spacing_before=0, spacing_after=60)

    # main.c с подробными комментариями
    main_commented = """\
#include "server.h"   /* start_server()      */
#include "config.h"   /* server_config_t, default_config(), load_config() */

/*
 * main — точка входа программы.
 *
 * Аргументы командной строки:
 *   argv[1] — необязательный путь к конфигурационному файлу.
 *              Если не указан, используются значения по умолчанию:
 *              address=0.0.0.0, port=8080, upload_dir=/tmp,
 *              username=admin, password=admin, auth_enabled=true.
 *
 * Возврат:
 *   Функция start_server() не завершается в штатном режиме (бесконечный
 *   цикл accept). Ненулевой код возврата означает ошибку инициализации.
 */
int main(int argc, char **argv)
{
    server_config_t cfg;

    /* Шаг 1: заполнить конфигурацию безопасными значениями по умолчанию */
    default_config(&cfg);

    /* Шаг 2: если передан файл конфигурации — загрузить его поверх дефолтов */
    if (argc > 1)
        load_config(argv[1], &cfg);

    /* Шаг 3: запустить TCP-сервер (функция не возвращается при нормальной работе) */
    return start_server(&cfg);
}"""

    for line in main_commented.splitlines():
        append(line, mono=True)

    # ── Приложение Б — server.c ───────────────────────────────────────────────
    append("", page_break=True)
    append("ПРИЛОЖЕНИЕ Б", bold=True, size=13,
           color=RGBColor(0x1D, 0x35, 0x57),
           align=WD_ALIGN_PARAGRAPH.CENTER,
           spacing_before=0, spacing_after=40)
    append("Листинг Б.1 — Реализация TCP-сервера (server.c)",
           bold=True, size=10,
           color=RGBColor(0x1D, 0x35, 0x57),
           align=WD_ALIGN_PARAGRAPH.CENTER,
           spacing_before=0, spacing_after=60)

    server_commented = """\
/*
 * server.c — главный модуль сервера.
 *
 * Реализует:
 *   - создание и настройку TCP-сокета (socket / bind / listen);
 *   - бесконечный цикл приёма соединений (accept);
 *   - чтение HTTP-заголовков (recv);
 *   - маршрутизацию запросов (GET / POST /upload);
 *   - проверку авторизации (check_auth);
 *   - вызов parse_multipart для извлечения файла из тела запроса;
 *   - журналирование событий (log_msg).
 *
 * Зависимости (только стандартные заголовки C11 + POSIX.1):
 *   <sys/socket.h>, <arpa/inet.h>, <unistd.h>, <sys/stat.h>,
 *   <sys/time.h>, <signal.h>, <errno.h>, <time.h>, <stdarg.h>
 */

""" + SRC["server.c"]

    for line in server_commented.splitlines():
        append(line, mono=True)

    doc.save(DOC)
    print(f"Saved → {DOC}")

if __name__ == "__main__":
    main()
