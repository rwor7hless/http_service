#!/usr/bin/env python3
"""
Insert flowchart PNGs into the report docx.
Does NOT touch any existing paragraph styles / fonts.
Run: python3 insert_diagrams.py
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from copy import deepcopy
import lxml.etree as etree

DOC_PATH = "Практическая_работа_HTTP_сервер.docx"
DIAG_DIR = "diagrams"

# ─────────────────────────────────────────────────────────────────────────────
# Step 1 — strip the code styling that was mistakenly applied
# ─────────────────────────────────────────────────────────────────────────────

def strip_code_styles(doc):
    """Remove Courier New font, grey shading, borders and left indent
    that were added to paragraphs by the previous script run."""
    fixed = 0
    for para in doc.paragraphs:
        changed = False

        # Fix runs: remove Courier New font override
        for run in para.runs:
            rPr = run._r.get_or_add_rPr()
            # remove rFonts if it forces Courier New
            for rf in rPr.findall(qn('w:rFonts')):
                ascii_val = rf.get(qn('w:ascii'), '')
                if 'Courier' in ascii_val:
                    rPr.remove(rf)
                    changed = True
            # remove sz (8.5pt = 170 half-points) that we injected
            for sz in rPr.findall(qn('w:sz')):
                if sz.get(qn('w:val')) == '170':
                    rPr.remove(sz)
                    changed = True
            for sz in rPr.findall(qn('w:szCs')):
                if sz.get(qn('w:val')) == '170':
                    rPr.remove(sz)
                    changed = True
            # remove colour 0D1B2A that we injected
            for co in rPr.findall(qn('w:color')):
                if co.get(qn('w:val'), '').upper() == '0D1B2A':
                    rPr.remove(co)
                    changed = True

        # Fix pPr: remove shading, borders, indent we injected
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            # shading with fill EDF2F4
            for shd in pPr.findall(qn('w:shd')):
                if shd.get(qn('w:fill'), '').upper() == 'EDF2F4':
                    pPr.remove(shd)
                    changed = True
            # paragraph border we added
            for pBdr in pPr.findall(qn('w:pBdr')):
                # check for our colour B0C4D8
                xml = etree.tostring(pBdr, encoding='unicode')
                if 'B0C4D8' in xml.upper():
                    pPr.remove(pBdr)
                    changed = True
            # indent left=360 we added
            for ind in pPr.findall(qn('w:ind')):
                if ind.get(qn('w:left')) == '360':
                    pPr.remove(ind)
                    changed = True
            # spacing before=40/after=40/line=240 we added
            for sp in pPr.findall(qn('w:spacing')):
                if sp.get(qn('w:before')) == '40' and sp.get(qn('w:line')) == '240':
                    pPr.remove(sp)
                    changed = True

        if changed:
            fixed += 1

    print(f"Stripped code styles from {fixed} paragraphs")


# ─────────────────────────────────────────────────────────────────────────────
# Step 2 — insert diagram images
# ─────────────────────────────────────────────────────────────────────────────

def para_index(doc, fragment):
    for i, p in enumerate(doc.paragraphs):
        if fragment in p.text:
            return i
    return -1


def add_image_block(doc, after_para_idx, img_path, caption, width_in=5.2):
    paras = doc.paragraphs
    ref_elem = paras[after_para_idx]._element

    # ── image paragraph ──
    img_p_el = OxmlElement('w:p')
    ref_elem.addnext(img_p_el)
    img_p = Paragraph(img_p_el, doc)
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(img_path, width=Inches(width_in))

    pPr = img_p_el.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '120')
    sp.set(qn('w:after'),  '40')
    pPr.append(sp)

    # ── caption paragraph ──
    cap_p_el = OxmlElement('w:p')
    img_p_el.addnext(cap_p_el)
    cap_p = Paragraph(cap_p_el, doc)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap_p.add_run(caption)
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x1D, 0x35, 0x57)

    pPr2 = cap_p_el.get_or_add_pPr()
    sp2 = OxmlElement('w:spacing')
    sp2.set(qn('w:before'), '0')
    sp2.set(qn('w:after'),  '160')
    pPr2.append(sp2)

    print(f"  Inserted after para {after_para_idx}: {caption[:60]}")


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

DIAGRAMS = [
    # Work bottom-up so earlier indices stay valid
    (
        'Функция default_config() инициализирует все поля',
        'diagram5_config.png',
        'Рисунок 5 — Алгоритм загрузки конфигурации (load_config / default_config)'
    ),
    (
        'После декодирования формируется ожидаемая строка',
        'diagram4_auth.png',
        'Рисунок 4 — Алгоритм проверки авторизации check_auth() и декодирования Base64'
    ),
    (
        'Использование memmem() вместо strstr()',
        'diagram3_multipart.png',
        'Рисунок 3 — Алгоритм парсинга тела запроса parse_multipart()'
    ),
    (
        'HTML-форма встроена в исходный код как строковая константа',
        'diagram2_request_routing.png',
        'Рисунок 2 — Алгоритм маршрутизации и обработки HTTP-запроса'
    ),
    (
        'Функция start_server() реализует основной цикл',
        'diagram1_server_loop.png',
        'Рисунок 1 — Главный цикл сервера (start_server)'
    ),
]

def already_has_images(doc):
    return len(doc.inline_shapes) > 0


def main():
    doc = Document(DOC_PATH)

    # Always strip mistakenly applied code styles first
    strip_code_styles(doc)

    # Only insert images if not already present
    if already_has_images(doc):
        print("Images already present — skipping insertion")
    else:
        for fragment, fname, caption in DIAGRAMS:
            idx = para_index(doc, fragment)
            if idx < 0:
                print(f"WARNING: anchor not found for {fname}")
                continue
            img_path = os.path.join(DIAG_DIR, fname)
            if not os.path.exists(img_path):
                print(f"WARNING: image not found: {img_path}")
                continue
            add_image_block(doc, idx, img_path, caption)

    doc.save(DOC_PATH)
    print(f"\nDone → {DOC_PATH}")


if __name__ == '__main__':
    main()
